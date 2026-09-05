"""Extract Word (.docx / converted .doc) into the same HTML shape as PDFs."""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from services.paginate import count_pages


def extract_word(path: str | Path, media_dir: Path) -> dict:
    path = Path(path)
    media_dir.mkdir(parents=True, exist_ok=True)
    work = path
    if path.suffix.lower() in {".doc", ".dot", ".rtf"}:
        work = _convert_legacy(path)
    html = _from_docx(work, media_dir)
    title = _title_from_docx(work) or path.stem
    pages = max(1, count_pages(html))
    tables = html.count("<table")
    return {
        "title": title,
        "page_count": pages,
        "table_count": tables,
        "pages": [],
        "html": html,
    }


def _from_docx(path: Path, media_dir: Path) -> str:
    html = _mammoth_html(path, media_dir)
    if html and ("<table" in html or "<p" in html or "<h" in html):
        return _wrap_pages(html, path.stem)
    return _wrap_pages(_python_docx_html(path, media_dir), path.stem)


def _mammoth_html(path: Path, media_dir: Path) -> str:
    try:
        import mammoth
    except ImportError:
        return ""

    counter = {"n": 0}

    def convert_image(image):
        counter["n"] += 1
        ext = (image.content_type or "image/png").split("/")[-1].replace("jpeg", "jpg")
        name = f"word_img{counter['n']}.{ext}"
        with image.open() as handle:
            (media_dir / name).write_bytes(handle.read())
        return {"src": f"__MEDIA__/{name}"}

    with path.open("rb") as handle:
        result = mammoth.convert_to_html(
            handle,
            convert_image=mammoth.images.img_element(convert_image),
        )
    raw = (result.value or "").strip()
    if not raw:
        return ""
    soup = BeautifulSoup(raw, "html.parser")
    for table in soup.find_all("table"):
        classes = table.get("class", [])
        if isinstance(classes, str):
            classes = classes.split()
        if "epdf-table" not in classes:
            table["class"] = " ".join(classes + ["epdf-table"]).strip()
    for img in soup.find_all("img"):
        wrap = soup.new_tag("figure")
        wrap["class"] = "epdf-figure"
        img.replace_with(wrap)
        wrap.append(img)
    return str(soup)


def _python_docx_html(path: Path, media_dir: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    img_i = 0
    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = (block.text or "").strip()
            style = (block.style.name or "") if block.style else ""
            if not text:
                continue
            esc = _esc(text)
            if style.startswith("Heading 1") or style == "Title":
                parts.append(f"<h1>{esc}</h1>")
            elif style.startswith("Heading"):
                parts.append(f"<h2>{esc}</h2>")
            else:
                parts.append(f"<p>{esc}</p>")
        elif isinstance(block, Table):
            rows = [[(cell.text or "").strip() for cell in row.cells] for row in block.rows]
            parts.append(_table_html(rows))
    for rel in doc.part.rels.values():
        try:
            reltype = getattr(rel, "reltype", "")
            if "image" not in reltype:
                continue
            blob = rel.target_part.blob
            ext = Path(str(getattr(rel, "target_ref", "png"))).suffix.lstrip(".") or "png"
        except Exception:
            continue
        img_i += 1
        name = f"word_rel{img_i}.{ext}"
        (media_dir / name).write_bytes(blob)
        parts.append(f"<figure class='epdf-figure'><img src='__MEDIA__/{name}' alt=''></figure>")
    return "".join(parts) or "<p></p>"


def _iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _table_html(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    head, body = rows[0], rows[1:]
    thead = "<thead><tr>" + "".join(f"<th>{_esc(c)}</th>" for c in head) + "</tr></thead>"
    tbody = "<tbody>"
    for row in body:
        tbody += "<tr>" + "".join(f"<td>{_esc(c)}</td>" for c in row) + "</tr>"
    tbody += "</tbody>"
    return f"<table class='epdf-table'>{thead}{tbody}</table>"


def _wrap_pages(inner: str, title: str) -> str:
    soup = BeautifulSoup(inner or "<p></p>", "html.parser")
    root = soup.body if soup.body else soup
    nodes = [n for n in list(root.children) if getattr(n, "name", None) or str(n).strip()]
    chunks: list[list] = [[]]
    for node in nodes:
        name = getattr(node, "name", "")
        if name in {"h1"} and chunks[-1]:
            chunks.append([node])
        else:
            chunks[-1].append(node)
    pages = []
    for i, chunk in enumerate(chunks, start=1):
        body = "".join(str(n) for n in chunk)
        pages.append(
            f"<section class='epdf-page' data-page='{i}'>{body}"
            f"<div class='epdf-page-footer'><span class='epdf-doc-title'>{_esc(title)}</span>"
            f"<span class='epdf-page-num'>Page {i}</span></div></section>"
        )
    return "".join(pages)


def _title_from_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        if doc.core_properties.title:
            return doc.core_properties.title
        for para in doc.paragraphs[:8]:
            if (para.text or "").strip():
                return para.text.strip()[:80]
    except Exception:
        return ""
    return ""


def _convert_legacy(path: Path) -> Path:
    dest = path.with_suffix(".docx")
    if dest.exists() and dest.stat().st_mtime >= path.stat().st_mtime:
        return dest
    err = _word_com(path, dest) or _libreoffice(path, dest)
    if dest.exists():
        return dest
    raise RuntimeError(
        err
        or "Old .doc/.rtf file convert nahi hua. Microsoft Word ya LibreOffice install karein, ya file ko .docx mein save karein."
    )


def _word_com(src: Path, dest: Path) -> str | None:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return "pywin32 missing"
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(src.resolve()))
        # 16 = wdFormatXMLDocument (.docx)
        doc.SaveAs2(str(dest.resolve()), FileFormat=16)
        doc.Close(False)
        return None
    except Exception as exc:
        return str(exc)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def _libreoffice(src: Path, dest: Path) -> str | None:
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(candidate).exists():
                soffice = candidate
                break
    if not soffice:
        return "LibreOffice not found"
    try:
        with tempfile.TemporaryDirectory() as tmp:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(src)],
                check=True,
                capture_output=True,
                timeout=120,
            )
            produced = next(Path(tmp).glob("*.docx"), None)
            if not produced:
                return "LibreOffice produced no docx"
            shutil.copyfile(produced, dest)
        return None
    except Exception as exc:
        return str(exc)


def _esc(text: str) -> str:
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
