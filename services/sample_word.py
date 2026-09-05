"""Build a sample .docx with a glossary table."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROWS = [
    ("Term", "Meaning"),
    ("CSR Rules", "Companies (CSR Policy) Rules, 2014."),
    ("CSRN", "CSR Network - internal implementation body."),
    ("FCRA", "Foreign Contribution (Regulation) Act, 2010."),
    ("GC", "General Circular issued by the MCA."),
    ("FY", "Financial year."),
    ("NGO", "Non-governmental organisation as implementing agency."),
    ("Schedule VII", "Activities listed in Schedule VII of the Companies Act, 2013."),
]


def build(path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "CSR Practical Handbook"

    h = doc.add_paragraph("CSR Practical Handbook")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(22)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph("Definitions & Abbreviations")
    sub.runs[0].font.size = Pt(16)

    doc.add_paragraph(
        "This sample Word file tests table extraction. Upload .docx or .doc — "
        "cells stay editable after the template is applied."
    )

    table = doc.add_table(rows=len(ROWS), cols=2)
    table.style = "Table Grid"
    for i, (term, meaning) in enumerate(ROWS):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = meaning
        if i == 0:
            for cell in table.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph("")
    h2 = doc.add_paragraph("How to use this template")
    h2.runs[0].bold = True
    h2.runs[0].font.size = Pt(16)
    doc.add_paragraph(
        "Choose a template, preview it, then open the editor. "
        "If a table does not extract, use Table tools to paste or convert text."
    )

    doc.save(path)
    return path
